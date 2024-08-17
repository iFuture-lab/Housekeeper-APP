from docx import Document
import os
from django.conf import settings

# Load the document


# Check if the file exists
template_path = 'contract_Recruitment.docx'
save_path = 'updated_document.docx'


# doc = Document(template_path)

replacements = {
    '{contract_number}': '123456',
    '{customer_id}': 'John',
    '{package_details}': 'Package A',
    '{payment_details}': 'Paid in full',
    '{contract_terms}': 'Standard Terms',
    '{start_date}': '2024-01-01',
    '{end_date}': '2024-12-31',
    '{status}': 'Active',
    '{request_type}': 'N'
}


def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    return doc

if not os.path.isfile(template_path):
    print(f"The file {template_path} does not exist.")
else:
    # Load the document
    doc = Document(template_path)

    # Replace placeholders
    doc = replace_placeholders(doc, replacements)

    # Save the updated document
    doc.save(save_path)
    print(f"Document saved successfully to {os.path.abspath(save_path)}")


