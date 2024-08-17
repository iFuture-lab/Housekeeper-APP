from docx import Document
import os
from django.conf import settings

# Load the document


# Check if the file exists
template_name = 'contract_Recruitment.docx'
template_path = os.path.join(settings.BASE_DIR, 'templates', template_name)

doc = Document(template_path)

# Define the replacements
replacements = {
    '{contract_number}': '123456',
    # Add other placeholders as needed
}

# Iterate through each paragraph in the document
for paragraph in doc.paragraphs:
    # To handle text within runs
    new_runs = []
    for run in paragraph.runs:
        # Replace placeholders in the run text
        new_text = run.text
        for key, value in replacements.items():
            if key in new_text:
                new_text = new_text.replace(key, value)
        if new_text != run.text:
            # Create a new run with the replaced text
            new_run = paragraph.add_run(new_text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_runs.append(new_run)
    # Remove old runs and replace with new runs
    for run in paragraph.runs:
        p = run._element
        p.getparent().remove(p)
    for new_run in new_runs:
        paragraph._element.append(new_run._element)

# Save the document with replacements
doc.save('updated_document.docx')
