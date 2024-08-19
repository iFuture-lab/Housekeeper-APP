from django.shortcuts import render
from rest_framework import generics
from rest_framework.response import Response
from .models import Contract
from .serializers import ContractSerializer
from docx import Document
import base64
import io
from drf_yasg.utils import swagger_auto_schema
from drf_yasg import openapi
from rest_framework.permissions import AllowAny
import os
from django.conf import settings
from docx import Document
from django.conf import settings
from django.http import HttpResponse
import os
from .models import UserInterest
from .serializers import UserInterestSerializer

from rest_framework.views import APIView
from rest_framework.response import Response
from django.db.models import Count
from rest_framework import status
from housekeeper.models import HireRequest,TransferRequest,RecruitmentRequest,Status
from service_type.models import ServiceType

class UserInterestCreateView(generics.CreateAPIView):
    queryset = UserInterest.objects.all()
    serializer_class = UserInterestSerializer
    permission_classes = [AllowAny]

class UserInterestListView(generics.ListAPIView):
    queryset = UserInterest.objects.all()
    serializer_class = UserInterestSerializer
    permission_classes = [AllowAny]
    

    def get_queryset(self):
        return UserInterest.objects.filter(user=self.request.user)
    
    
class UserInterestReportView(APIView):
    permission_classes = [AllowAny]

    def get(self, request, *args, **kwargs):
        data = UserInterest.objects.values('status').annotate(total=Count('status'))
        return Response(data)
    
    
####################contract#####################################

class ContractCreateView(generics.CreateAPIView):
    permission_classes = [AllowAny]
    queryset = Contract.objects.all()
    serializer_class = ContractSerializer
      
    def perform_create(self, serializer):
        contract = serializer.save(status='Pending')
        request_type = contract.request_type
        print("'''''''''''''''''",request_type)
        print("Request Type Attributes:", dir(request_type))
    
    # Attempt to retrieve possible name attributes
        request_type_name = getattr(request_type, 'name', None)
        print("Request Type Name:", request_type_name)
        if request_type:
            print("""""""""""""""""","hi")
        x= ServiceType.objects.get(name='Hire')
        y= ServiceType.objects.get(name='Transfer')
        z= ServiceType.objects.get(name='Recruitment')
        
        if request_type == x:
            print("I got ittttttttttttttttttt")  
            valid_payment = self.check_hire_request_status(contract)
            print("found hire request paid",valid_payment)
        elif request_type == y:
            valid_payment = self.check_transfer_request_status(contract)
        elif request_type == z:
            valid_payment = self.check_recruitment_request_status(contract)
        else:
            valid_payment = False
            
        if valid_payment:
            try:
                # encoded_contract = self.generate_contract(contract)
                # contract.contract_file = encoded_contract
                # contract.status = 'Completed'
                # contract.save()
                
                # Update the status first
                contract.status = 'Completed'
                contract.save(update_fields=['status'])
            
            # Then generate the contract
                encoded_contract = self.generate_contract(contract)
                contract.contract_file = encoded_contract
                contract.save(update_fields=['contract_file'])
                return Response({
                    'contract_number': contract.contract_number,
                    'contract_file': encoded_contract
                }, status=status.HTTP_201_CREATED)

            except Exception as e:
                return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        else:
           
            contract.delete()
            return Response({'error': 'Payment status is not valid.'}, status=status.HTTP_400_BAD_REQUEST)


      

    def check_hire_request_status(self, contract):
        if contract.hire_request:
            hire_request = HireRequest.objects.filter(
                id=contract.hire_request.id).first()
            status= Status.objects.get(Status='Paid')
            if hire_request.status== status:
                print("''''''''''''","I found itttttttttttt")
                return hire_request
        return False

    
    def check_transfer_request_status(self, contract):
        if contract.transfer_request:
            transfer_request = TransferRequest.objects.filter(
                id=contract.transfer_request.id
            ).first()
            status= Status.objects.get(Status='Paid')
            if transfer_request.status== status:
                print("''''''''''''","I found itttttttttttt")
            return transfer_request
        return False
    
    
    def check_recruitment_request_status(self, contract):
        if contract.recruitment_request:
            recruitment_request = RecruitmentRequest.objects.filter(
                id=contract.recruitment_request.id
            ).first()
            status= Status.objects.get(Status='Paid')
            if recruitment_request.status== status:
                print("''''''''''''","I found itttttttttttt")
            return recruitment_request
        return False
    
        
        
         
    def check_template_path(request):
        template_name = 'contract_Recruitment.docx'
        template_path = os.path.join(settings.BASE_DIR, 'templates', template_name)
        print(template_path)
        if os.path.exists(template_path):
            print("hhhhhhhhhhhhhhhhhhhhhhh")
            return HttpResponse(f'Template found at: {template_path}')
        else:
            print("notttttttttttttttttttttttttt")
            return HttpResponse(f'Template not found at: {template_path}')
        
        
    def generate_contract(self, contract):
        template_name = 'contract_Recruitment.docx'
        template_path = os.path.join(settings.BASE_DIR, 'templates', template_name)
        save_path = 'contract.docx'

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found at {template_path}")

 
        replacements = {
            '{contract_number}': str(contract.contract_number) if contract.contract_number else '',
            '{customer_id}': str(contract.customer_id.fullName) if contract.customer_id else '',
            '{package_details}': str(contract.package_details) if contract.package_details else '',
            '{payment_details}': str(contract.payment_details.order_id) if contract.payment_details else '',
            '{contract_terms}': str(contract.contract_terms) if contract.contract_terms else '',
            '{start_date}': contract.start_date.strftime('%Y/%m/%d') if contract.start_date else '',
            '{end_date}': contract.end_date.strftime('%Y/%m/%d') if contract.end_date else '',
            '{status}': contract.status if contract.status else '',
            '{request_type}': str(contract.request_type) if contract.request_type else '',
        }

        def replace_placeholders(doc, replacements):
            for paragraph in doc.paragraphs:
               
                original_text = paragraph.text
                print(f"Original text in paragraph: '{original_text}'")
        
           
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
                        print(f"Replaced '{key}' with '{value}'")
                updated_text = paragraph.text
                if original_text != updated_text:
                    print(f"Updated text in paragraph: '{updated_text}'")
            return doc


        doc = Document(template_path)
        print(f"Document loaded from {template_path}")

        doc = replace_placeholders(doc, replacements)
          
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        doc.save(save_path)

        encoded_file = base64.b64encode(buffer.read()).decode('utf-8')
        return encoded_file
        

        
        # print(f"Document saved successfully to {os.path.abspath(save_path)}")
        


class ContractListView(generics.ListAPIView):
    permission_classes = [AllowAny]
    queryset = Contract.objects.all()
    serializer_class = ContractSerializer

class ContractDetailView(generics.RetrieveUpdateDestroyAPIView):
    permission_classes = [AllowAny]
    queryset = Contract.objects.all()
    serializer_class = ContractSerializer
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
 # Save the updated document
    
        # buffer = io.BytesIO()
        # doc.save(buffer)
        # buffer.seek(0)

        # encoded_file = base64.b64encode(buffer.read()).decode('utf-8')
        # return encoded_file
        # doc.save()
        # print(f"Document saved successfully to {os.path.abspath(save_path)}")
        
        
        
        
        
        
                  
    
    

 

    

    # def perform_create(self, serializer):
    #     contract = serializer.save(status='Pending')
    #     encoded_contract = self.generate_contract(contract)
    #     contract.contract_file = encoded_contract
    #     contract.status = 'Completed'
    #     contract.save()
        
    

    # def generate_contract(self, contract):
    #     template_name = 'contract_Recruitment.docx'
    #     template_path = os.path.join(settings.BASE_DIR, 'templates', template_name)
        
    #     print(f"Template path: {template_path}")
    #     print(f"File exists: {os.path.exists(template_path)}")
        
    #     if not os.path.exists(template_path):
    #         raise FileNotFoundError(f"Template file not found at {template_path}")
        
    #     print(template_path)
        
        
    #     doc = Document(template_path)
    #     print(doc)
    
        
        # for paragraph in doc.paragraphs:
        #      paragraph.text = paragraph.text.replace('{hi}', contract.contract_number)
            # paragraph.text = paragraph.text.replace('{customer_id}', str(contract.customer_id.Fullname))
            # paragraph.text = paragraph.text.replace('{package_details}', contract.package_details)
            # paragraph.text = paragraph.text.replace('{payment_details}', contract.payment_details)
            # paragraph.text = paragraph.text.replace('{contract_terms}', contract.contract_terms)
            # paragraph.text = paragraph.text.replace('{start_date}', str(contract.start_date))
            # paragraph.text = paragraph.text.replace('{end_date}', str(contract.end_date))
            # paragraph.text = paragraph.text.replace('{status}', contract.status)
            # paragraph.text = paragraph.text.replace('{request_type}', contract.request_type)

       
        # buffer = io.BytesIO()
        # doc.save(buffer)
        # buffer.seek(0)

        # encoded_file = base64.b64encode(buffer.read()).decode('utf-8')
        # return encoded_file
        
        
        # contract = serializer.save(status='Pending')

        # # Generate and encode the contract
        # try:
        #     encoded_contract = self.generate_contract(contract)
        #     contract.contract_file = encoded_contract
        #     contract.status = 'Completed'
        #     contract.save()
        #     return Response({
        #     'contract_number': contract.contract_number,
        #     'contract_file': encoded_contract
        # }, status=status.HTTP_201_CREATED)
            
        # except Exception as e:
        #     # Handle errors gracefully and provide feedback
        #     return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        
        
